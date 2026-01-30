from __future__ import annotations

import argparse
import io
import json
import os
import tempfile
from pathlib import Path
from typing import List

import time
from tqdm import tqdm

from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings

from configuracao import carregar_propriedades
from drive_api import (
    baixar_arquivo,
    exportar_texto_plano,
    extrair_id_pasta,
    listar_arquivos_pasta,
)
from texto_utils import limpar_texto, separar_secoes, fatiar_secoes
from observabilidade import registrar_evento


def extrair_texto_docx(caminho: Path) -> str:
    from docx import Document

    doc = Document(str(caminho))
    return "\n".join(p.text for p in doc.paragraphs)


def extrair_texto_docx_bytes(conteudo: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(conteudo))
    return "\n".join(p.text for p in doc.paragraphs)


def extrair_texto_html(conteudo: str) -> str:
    try:
        from bs4 import BeautifulSoup
    except Exception:
        import re

        texto = re.sub(r"<[^>]+>", " ", conteudo)
        return " ".join(texto.split())

    soup = BeautifulSoup(conteudo, "html.parser")
    return " ".join(soup.get_text(" ").split())


def configurar_ambiente_textract() -> None:
    caminhos_bins = [
        r"C:\msys64\mingw64\bin",
        r"C:\msys64\ucrt64\bin",
        r"C:\msys64\usr\bin",
    ]
    path_atual = os.environ.get("PATH", "")
    for caminho_bin in caminhos_bins:
        if caminho_bin not in path_atual:
            path_atual = f"{caminho_bin};{path_atual}"
    os.environ["PATH"] = path_atual
    if not os.environ.get("HOME"):
        os.environ["HOME"] = os.environ.get("USERPROFILE", "")


def extrair_texto_textract_bytes(conteudo: bytes, sufixo: str) -> str:
    if conteudo.startswith(b"PK\x03\x04"):
        return extrair_texto_docx_bytes(conteudo)
    try:
        import textract
    except Exception as exc:  # pragma: no cover - dependency error
        raise RuntimeError("textract nao instalado. Instale requirements.txt.") from exc

    configurar_ambiente_textract()

    with tempfile.NamedTemporaryFile(suffix=sufixo, delete=False) as tmp:
        tmp.write(conteudo)
        tmp_path = Path(tmp.name)
    try:
        resultado = textract.process(str(tmp_path))
        return resultado.decode("utf-8", errors="ignore")
    except Exception as exc:  # pragma: no cover - external tool error
        try:
            texto = conteudo.decode("utf-8", errors="ignore")
            if "<html" in texto.lower():
                return extrair_texto_html(texto)
        except Exception:
            pass
        raise RuntimeError(f"Falha ao ler arquivo com textract: {exc}") from exc
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass


def extrair_texto_textract_arquivo(caminho: Path) -> str:
    try:
        import textract
    except Exception as exc:  # pragma: no cover - dependency error
        raise RuntimeError("textract nao instalado. Instale requirements.txt.") from exc

    configurar_ambiente_textract()
    try:
        resultado = textract.process(str(caminho))
        return resultado.decode("utf-8", errors="ignore")
    except Exception as exc:  # pragma: no cover - external tool error
        try:
            texto = caminho.read_text(encoding="utf-8", errors="ignore")
            if "<html" in texto.lower():
                return extrair_texto_html(texto)
        except Exception:
            pass
        raise RuntimeError(f"Falha ao ler arquivo com textract: {exc}") from exc



def extrair_texto(caminho: Path) -> str:
    sufixo = caminho.suffix.lower()
    if sufixo == ".docx":
        return extrair_texto_docx(caminho)

    if sufixo == ".doc":
        return extrair_texto_textract_arquivo(caminho)

    raise RuntimeError("Formato nao suportado. Use .docx ou .doc.")


def extrair_texto_doc_bytes(conteudo: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(conteudo)
        tmp_path = Path(tmp.name)
    try:
        return extrair_texto(tmp_path)
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass


def iterar_documentos(pasta_entrada: Path) -> List[Path]:
    return sorted(
        [p for p in pasta_entrada.rglob("*") if p.suffix.lower() in {".doc", ".docx"}]
    )


def criar_indice(
    pasta_entrada: Path,
    pasta_indice: Path,
    modelo: str,
    max_caracteres: int,
    sobreposicao: int,
) -> None:
    pasta_entrada = pasta_entrada.resolve()
    pasta_indice.mkdir(parents=True, exist_ok=True)

    modelo_embeddings = HuggingFaceEmbeddings(model_name=modelo)
    todos_chunks = []
    falhas = []
    inicio = time.time()
    registrar_evento(
        "index_inicio",
        pasta_entrada="drive",
        pasta_indice=str(pasta_indice),
        modelo=modelo,
        max_caracteres=max_caracteres,
        sobreposicao=sobreposicao,
    )

    for caminho in tqdm(iterar_documentos(pasta_entrada), desc="Lendo documentos"):
        try:
            texto = limpar_texto(extrair_texto(caminho))
        except Exception as exc:
            falhas.append({"arquivo": caminho.name, "erro": str(exc)})
            continue
        secoes = separar_secoes(texto)
        chunks = fatiar_secoes(
            secoes, max_caracteres=max_caracteres, sobreposicao=sobreposicao
        )
        for i, chunk in enumerate(chunks):
            todos_chunks.append(
                {
                    "id": f"{caminho.stem}-{i}",
                    "file_name": caminho.name,
                    "title": chunk["title"],
                    "text": chunk["text"],
                }
            )

    if not todos_chunks:
        registrar_evento("index_vazio", pasta_entrada=str(pasta_entrada))
        if falhas:
            registrar_evento(
                "index_falhas",
                pasta_entrada=str(pasta_entrada),
                falhas=falhas,
            )
        raise RuntimeError("Nenhum documento indexado. Verifique erros de leitura.")

    textos = [c["text"] for c in todos_chunks]
    metadados = [
        {"arquivo": c["file_name"], "titulo": c["title"], "id": c["id"]}
        for c in todos_chunks
    ]

    indice = FAISS.from_texts(textos, modelo_embeddings, metadatas=metadados)
    indice.save_local(str(pasta_indice))

    with open(pasta_indice / "config.json", "w", encoding="utf-8") as f:
        json.dump(
            {
                "modelo": modelo,
                "max_caracteres": max_caracteres,
                "sobreposicao": sobreposicao,
                "total_trechos": len(todos_chunks),
            },
            f,
            ensure_ascii=True,
            indent=2,
        )

    registrar_evento(
        "index_fim",
        pasta_entrada=str(pasta_entrada),
        pasta_indice=str(pasta_indice),
        modelo=modelo,
        total_trechos=len(todos_chunks),
        duracao_seg=round(time.time() - inicio, 3),
    )
    if falhas:
        registrar_evento(
            "index_falhas",
            pasta_entrada=str(pasta_entrada),
            falhas=falhas,
        )


def ler_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Indexador RAG simples")
    parser.add_argument("--input", default="data/raw", help="Pasta com .doc/.docx")
    parser.add_argument("--index-dir", default="index", help="Pasta do indice")
    parser.add_argument(
        "--model",
        default="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
        help="Modelo de embeddings",
    )
    parser.add_argument("--max-caracteres", type=int, default=1200)
    parser.add_argument("--sobreposicao", type=int, default=200)
    return parser.parse_args()


def main() -> None:
    args = ler_args()
    criar_indice(
        Path(args.input),
        Path(args.index_dir),
        args.model,
        args.max_caracteres,
        args.sobreposicao,
    )


if __name__ == "__main__":
    main()
